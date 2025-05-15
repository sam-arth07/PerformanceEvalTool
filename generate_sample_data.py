#!/usr/bin/env python
"""
Demo Data Generator for PET Application
--------------------------------------

This script generates sample resumes and video files for testing the PET application.
"""

import logging
import os
import random
import shutil
import string
from datetime import datetime, timedelta
from pathlib import Path

from docx import Document

# Set up logging
logging.basicConfig(level=logging.INFO, format='%(asctime)s - %(levelname)s - %(message)s')
logger = logging.getLogger(__name__)

def create_sample_directory():
    """Create a directory for sample files"""
    sample_dir = Path("sample_data")
    sample_dir.mkdir(exist_ok=True)
    
    resume_dir = sample_dir / "resumes"
    resume_dir.mkdir(exist_ok=True)
    
    video_dir = sample_dir / "videos"
    video_dir.mkdir(exist_ok=True)
    
    logger.info(f"Created sample directories: {sample_dir}")
    
    return resume_dir, video_dir

def generate_resume_text(candidate_level='mid'):
    """Generate sample resume text based on candidate level"""
    
    # Basic candidate information
    first_names = ["John", "Jane", "Michael", "Sarah", "David", "Emily", "Robert", "Lisa"]
    last_names = ["Smith", "Johnson", "Williams", "Brown", "Jones", "Miller", "Davis", "Wilson"]
    
    name = f"{random.choice(first_names)} {random.choice(last_names)}"
    email = f"{name.lower().replace(' ', '.')}@example.com"
    phone = f"555-{random.randint(100, 999)}-{random.randint(1000, 9999)}"
    
    # Education section
    universities = [
        "University of Technology", "State University", 
        "National Institute of Science", "Metropolitan University"
    ]
    
    degrees = {
        'entry': ["Bachelor of Science", "Bachelor of Arts", "Associate Degree"],
        'mid': ["Bachelor of Science", "Master of Science", "Bachelor of Technology"],
        'senior': ["Master of Science", "Master of Engineering", "Ph.D."]
    }
    
    majors = {
        'entry': ["Computer Science", "Information Technology", "Software Engineering"],
        'mid': ["Computer Science", "Software Engineering", "Data Science"],
        'senior': ["Computer Science", "Artificial Intelligence", "Machine Learning"]
    }
    
    # Get current year and calculate graduation year
    current_year = datetime.now().year
    grad_years = {
        'entry': list(range(current_year-3, current_year+1)),
        'mid': list(range(current_year-7, current_year-2)),
        'senior': list(range(current_year-15, current_year-5))
    }
    
    university = random.choice(universities)
    degree = random.choice(degrees[candidate_level])
    major = random.choice(majors[candidate_level])
    grad_year = random.choice(grad_years[candidate_level])
    
    # Skills section
    technical_skills = {
        'entry': [
            "HTML", "CSS", "JavaScript", "Python basics", 
            "SQL basics", "Git", "Microsoft Office"
        ],
        'mid': [
            "Java", "Python", "JavaScript", "React", "Node.js",
            "SQL", "MongoDB", "Docker basics", "CI/CD", "Git"
        ],
        'senior': [
            "Java", "Python", "Kotlin", "JavaScript", "TypeScript",
            "React", "Angular", "Node.js", "AWS", "Azure", 
            "Docker", "Kubernetes", "CI/CD", "Microservices"
        ]
    }
    
    # Experience section
    companies = [
        "Tech Solutions Inc.", "Innovate Systems", "Digital Dynamics",
        "CodeCraft Technologies", "Future Software Ltd.", "Next Level Tech"
    ]
    
    junior_titles = ["Junior Developer", "Software Intern", "Junior Software Engineer", "Entry Level Programmer"]
    mid_titles = ["Software Developer", "Software Engineer", "Full Stack Developer", "Application Developer"]
    senior_titles = ["Senior Software Engineer", "Lead Developer", "Software Architect", "Development Manager"]
    
    experience_years = {
        'entry': list(range(0, 3)),
        'mid': list(range(3, 8)),
        'senior': list(range(8, 20))
    }
    
    exp_years = random.choice(experience_years[candidate_level])
    
    # Generate job history based on experience level
    job_history = []
    
    if candidate_level == 'entry':
        # 0-2 jobs
        num_jobs = random.randint(0, 2)
        title_list = junior_titles
    elif candidate_level == 'mid':
        # 1-3 jobs
        num_jobs = random.randint(1, 3)
        title_list = mid_titles
    else:  # senior
        # 2-4 jobs
        num_jobs = random.randint(2, 4)
        title_list = senior_titles
    
    # Create job history
    end_date = datetime.now()
    for i in range(num_jobs):
        company = random.choice(companies)
        title = random.choice(title_list)
        
        # Calculate duration
        if i == 0:  # Current job
            duration = random.randint(1, 36)  # 1-36 months
            start_date = end_date - timedelta(days=duration*30)
            date_str = f"{start_date.strftime('%b %Y')} - Present"
        else:
            duration = random.randint(6, 30)  # 6-30 months
            start_date = end_date - timedelta(days=duration*30)
            date_str = f"{start_date.strftime('%b %Y')} - {end_date.strftime('%b %Y')}"
            end_date = start_date - timedelta(days=random.randint(15, 60))  # Gap between jobs
        
        # Job duties
        if 'Junior' in title or 'Intern' in title:
            duties = [
                "Assisted in developing software components under supervision",
                "Fixed bugs and performed testing tasks",
                "Worked on UI implementation following designs",
                "Documented code and technical processes"
            ]
        elif 'Senior' in title or 'Lead' in title or 'Architect' in title:
            duties = [
                "Led team of developers in delivering critical projects",
                "Designed and implemented system architecture",
                "Mentored junior developers and conducted code reviews",
                "Collaborated with stakeholders to define technical requirements",
                "Implemented DevOps practices and CI/CD pipelines"
            ]
        else:
            duties = [
                "Developed and maintained software applications",
                "Implemented new features and functionality",
                "Collaborated with cross-functional teams",
                "Participated in code reviews and testing",
                "Fixed bugs and improved application performance"
            ]
        
        selected_duties = random.sample(duties, k=min(3, len(duties)))
        
        job_history.append({
            "company": company,
            "title": title,
            "date": date_str,
            "duties": selected_duties
        })
    
    # Build the resume text
    resume_sections = []
    
    # Contact info
    resume_sections.append(f"{name}\n{email} | {phone}\n")
    
    # Summary section
    if candidate_level == 'entry':
        summary = f"Recent {degree} graduate in {major} with {exp_years} years of experience seeking an entry-level software development position to apply academic knowledge in a professional environment."
    elif candidate_level == 'mid':
        summary = f"Experienced software developer with {exp_years} years of experience in building and maintaining applications. Proficient in multiple programming languages and technologies with a strong focus on delivering quality code."
    else:  # senior
        summary = f"Senior software professional with {exp_years} years of industry experience leading teams and architecting complex systems. Proven track record of delivering high-quality software solutions and mentoring junior developers."
    
    resume_sections.append(f"PROFESSIONAL SUMMARY\n{summary}\n")
    
    # Education section
    resume_sections.append("EDUCATION")
    resume_sections.append(f"{university}")
    resume_sections.append(f"{degree} in {major}, {grad_year}\n")
    
    # Skills section
    resume_sections.append("TECHNICAL SKILLS")
    selected_skills = random.sample(technical_skills[candidate_level], 
                                   k=min(7, len(technical_skills[candidate_level])))
    resume_sections.append(", ".join(selected_skills) + "\n")
    
    # Experience section
    if job_history:
        resume_sections.append("PROFESSIONAL EXPERIENCE")
        for job in job_history:
            resume_sections.append(f"{job['title']} | {job['company']}")
            resume_sections.append(job['date'])
            for duty in job['duties']:
                resume_sections.append(f"• {duty}")
            resume_sections.append("")  # Empty line between jobs
    
    # Projects section if entry level
    if candidate_level == 'entry':
        resume_sections.append("PROJECTS")
        project_names = ["E-commerce Website", "Task Manager App", "Data Analysis Tool"]
        project = random.choice(project_names)
        resume_sections.append(f"{project}")
        resume_sections.append("• Developed using Python and JavaScript")
        resume_sections.append("• Implemented user authentication and data storage")
        resume_sections.append("• Created responsive UI/UX design\n")
    
    return "\n".join(resume_sections)

def create_sample_resume(resume_dir, candidate_level):
    """Create a sample resume DOCX file"""
    resume_text = generate_resume_text(candidate_level)
    
    # Create a new Document
    doc = Document()
    
    # Add the text to the document
    for line in resume_text.split('\n'):
        if line.strip() and any(section in line for section in ["SUMMARY", "EDUCATION", "SKILLS", "EXPERIENCE", "PROJECTS"]):
            # Add as heading
            doc.add_heading(line, level=1)
        else:
            # Add as paragraph
            doc.add_paragraph(line)
    
    # Save the document
    filename = f"sample_resume_{candidate_level}_{random.randint(1000, 9999)}.docx"
    file_path = resume_dir / filename
    doc.save(file_path)
    
    logger.info(f"Created sample resume: {file_path}")
    return file_path

def create_empty_video_file(video_dir, candidate_level):
    """
    Create an empty video file for testing
    
    In a real implementation, this would generate actual video content,
    but for testing purposes we'll just create placeholder files
    """
    filename = f"sample_video_{candidate_level}_{random.randint(1000, 9999)}.mp4"
    file_path = video_dir / filename
    
    # Create an empty file
    with open(file_path, "wb") as f:
        # Write a small amount of data so it's not completely empty
        f.write(b"\x00" * 1024)
    
    logger.info(f"Created sample video placeholder: {file_path}")
    return file_path

def main():
    """Main function to generate sample data"""
    logger.info("Generating sample data for PET application...")
    
    # Create sample directories
    resume_dir, video_dir = create_sample_directory()
    
    # Generate samples for each candidate level
    levels = ['entry', 'mid', 'senior']
    samples_per_level = 3
    
    for level in levels:
        logger.info(f"Generating {samples_per_level} samples for {level}-level candidates...")
        for i in range(samples_per_level):
            resume_path = create_sample_resume(resume_dir, level)
            video_path = create_empty_video_file(video_dir, level)
    
    logger.info("Sample data generation complete!")
    logger.info(f"Sample resumes: {resume_dir}")
    logger.info(f"Sample videos: {video_dir}")

if __name__ == "__main__":
    main()
